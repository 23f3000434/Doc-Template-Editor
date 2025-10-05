from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Load your WCR document
doc = Document('WCR.docx')  # Ya jo bhi filename hai

# Access Table 0, Row 3, Cell 2 (address wali cell)
cell = doc.tables[0].rows[3].cells[2]

# Clear any existing content
for para in cell.paragraphs:
    for run in para.runs:
        run.text = ''

# Add "Addressvariable" text with yellow highlight
para = cell.paragraphs[0]  # First paragraph in cell
run = para.add_run('Addressvariable')
run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# Save the document
doc.save('WCR-2-FIXED.docx')
print("✓ Done! Check WCR-2-FIXED.docx")